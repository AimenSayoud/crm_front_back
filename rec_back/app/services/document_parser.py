import PyPDF2
import docx
import pdfplumber
import logging
from typing import Optional, Union
from pathlib import Path
import io

logger = logging.getLogger(__name__)

class DocumentParser:
    """Service for parsing different document formats and extracting text"""
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF using both PyPDF2 and pdfplumber for better accuracy
        
        Args:
            file_content: PDF file content as bytes
            
        Returns:
            Extracted text as string
            
        Raises:
            ValueError: If text extraction fails
        """
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                
                if text.strip():
                    logger.info(f"Successfully extracted {len(text)} characters using pdfplumber")
                    return text.strip()
            
            # Fallback to PyPDF2 if pdfplumber fails
            logger.warning("pdfplumber extraction failed, trying PyPDF2")
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            if text.strip():
                logger.info(f"Successfully extracted {len(text)} characters using PyPDF2")
                return text.strip()
            else:
                raise ValueError("No text could be extracted from PDF")
                
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise ValueError(f"Failed to extract text from PDF file: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_content: DOCX file content as bytes
            
        Returns:
            Extracted text as string
            
        Raises:
            ValueError: If text extraction fails
        """
        try:
            doc = docx.Document(io.BytesIO(file_content))
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            # Extract text from headers and footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + "\n"
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + "\n"
            
            if text.strip():
                logger.info(f"Successfully extracted {len(text)} characters from DOCX")
                return text.strip()
            else:
                raise ValueError("No text could be extracted from DOCX")
                
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise ValueError(f"Failed to extract text from DOCX file: {str(e)}")
    
    @staticmethod
    def extract_text_from_file(file_content: bytes, content_type: str, filename: str = "") -> str:
        """
        Main method to extract text based on file type
        
        Args:
            file_content: File content as bytes
            content_type: MIME type of the file
            filename: Original filename for additional type detection
            
        Returns:
            Extracted text as string
            
        Raises:
            ValueError: If file type is unsupported or extraction fails
        """
        logger.info(f"Extracting text from file type: {content_type}, filename: {filename}")
        
        # PDF files
        if content_type in ['application/pdf'] or filename.lower().endswith('.pdf'):
            return DocumentParser.extract_text_from_pdf(file_content)
        
        # DOCX files
        elif (content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'] or 
              filename.lower().endswith('.docx')):
            return DocumentParser.extract_text_from_docx(file_content)
        
        # Plain text files
        elif content_type in ['text/plain'] or filename.lower().endswith(('.txt', '.md')):
            try:
                text = file_content.decode('utf-8')
                logger.info(f"Successfully decoded text file with {len(text)} characters")
                return text
            except UnicodeDecodeError:
                # Try other encodings
                for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        text = file_content.decode(encoding)
                        logger.info(f"Successfully decoded text file with encoding {encoding}")
                        return text
                    except UnicodeDecodeError:
                        continue
                raise ValueError("Could not decode text file with supported encodings")
        
        else:
            raise ValueError(f"Unsupported file type: {content_type} (filename: {filename})")
    
    @staticmethod
    def validate_file_type(content_type: str, filename: str = "") -> bool:
        """
        Validate if file type is supported
        
        Args:
            content_type: MIME type of the file
            filename: Original filename
            
        Returns:
            True if file type is supported, False otherwise
        """
        supported_types = {
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'text/plain'
        }
        
        supported_extensions = {'.pdf', '.docx', '.txt', '.md'}
        
        # Check MIME type
        if content_type in supported_types:
            return True
        
        # Check file extension as fallback
        if filename:
            file_ext = Path(filename).suffix.lower()
            if file_ext in supported_extensions:
                return True
        
        return False
    
    @staticmethod
    def get_file_info(file_content: bytes, content_type: str, filename: str = "") -> dict:
        """
        Get information about the uploaded file
        
        Args:
            file_content: File content as bytes
            content_type: MIME type of the file
            filename: Original filename
            
        Returns:
            Dictionary with file information
        """
        return {
            "filename": filename,
            "content_type": content_type,
            "size_bytes": len(file_content),
            "size_kb": round(len(file_content) / 1024, 2),
            "size_mb": round(len(file_content) / (1024 * 1024), 2),
            "is_supported": DocumentParser.validate_file_type(content_type, filename),
            "file_extension": Path(filename).suffix.lower() if filename else ""
        }